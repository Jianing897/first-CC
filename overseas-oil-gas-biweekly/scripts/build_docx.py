#!/usr/bin/env python3
"""将结构化 JSON 生成为《海外油气投资环境资讯》双周报 Word 文档。

版式严格对齐 references/sample_issue60.docx（第6期样例）。
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "references" / "sample_issue60.docx"
SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from coverage_rules import ISSUE_ARTICLE_COUNT, MAX_BODY_CHARS, count_chinese_chars

# 样例模板字体与缩进（EMU，与第6期 docx 一致）
FONT_TABLE_TITLE = "黑体"
FONT_SECTION = "方正楷体简体"
FONT_BODY = "仿宋"
FONT_LATIN = "Times New Roman"
INDENT_TITLE_LEFT = 266700
INDENT_BODY_FIRST = 304800
TABLE_FONT_COLOR = RGBColor(0, 0, 253)
LINE_SPACING_BODY = 1.5


def _ensure_rfonts(run):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    return rfonts


def _set_run_font(
    run,
    east_asia: str,
    size_pt: float = 12,
    bold: bool | None = None,
    latin: str = FONT_LATIN,
    color: RGBColor | None = None,
):
    run.font.name = latin
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    if color is not None:
        color_el = rpr.find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            rpr.append(color_el)
        color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _resolve_style(doc: Document, style_name: str) -> str:
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return "Normal"


def _set_paragraph_format(
    paragraph,
    *,
    align=None,
    left_indent=None,
    first_line_indent=None,
    line_spacing: float | None = None,
):
    pf = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.line_spacing = line_spacing


def _add_run_paragraph(
    doc: Document,
    text: str,
    *,
    style: str = "Normal",
    east_asia: str = FONT_BODY,
    size_pt: float = 12,
    bold: bool | None = None,
    color: RGBColor | None = None,
    align=None,
    left_indent=None,
    first_line_indent=None,
    line_spacing: float | None = None,
):
    resolved = _resolve_style(doc, style)
    paragraph = doc.add_paragraph(style=resolved)
    _set_paragraph_format(
        paragraph,
        align=align,
        left_indent=left_indent,
        first_line_indent=first_line_indent,
        line_spacing=line_spacing,
    )
    run = paragraph.add_run(text)
    _set_run_font(run, east_asia, size_pt, bold, color=color)
    return paragraph


def _clear_body_keep_header_table(doc: Document):
    body = doc.element.body
    children = list(body)
    for child in children[1:]:
        body.remove(child)


def _split_org(org: str) -> tuple[str, str]:
    if " " in org:
        return org.split(" ", 1)
    marker = "中国海洋石油集团能源经济研究院"
    if org.startswith(marker):
        rest = org[len(marker) :].strip()
        return marker, rest
    return org, ""


def _clear_paragraph_runs(paragraph):
    p_el = paragraph._element
    for child in list(p_el):
        if child.tag == qn("w:r"):
            p_el.remove(child)


def _write_cell_lines(cell, lines: list[str], *, center: bool = False, size_pt: float = 12, bold: bool | None = None):
    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph()
    for idx, text in enumerate(lines):
        paragraph = cell.paragraphs[idx]
        _clear_paragraph_runs(paragraph)
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_format(paragraph, line_spacing=LINE_SPACING_BODY)
        run = paragraph.add_run(text)
        _set_run_font(run, FONT_TABLE_TITLE, size_pt, bold, color=TABLE_FONT_COLOR)


def _fill_header_table(doc: Document, meta: dict):
    if not doc.tables:
        return
    table = doc.tables[0]
    title = meta.get("title", "海外油气投资环境资讯")
    subtitle = meta.get("subtitle", "（双周报）")
    org_line1, org_line2 = _split_org(meta["org"])
    issue_label = f" {meta['year']}年第{meta['issue']}期（总第{meta['total_issue']}期）"
    date_range = meta["date_range"]

    _write_cell_lines(table.rows[0].cells[0], [title], center=True, size_pt=36, bold=True)
    _write_cell_lines(table.rows[0].cells[1], [title], center=True, size_pt=36, bold=True)
    _write_cell_lines(table.rows[1].cells[0], [subtitle], center=True, size_pt=16)
    _write_cell_lines(table.rows[1].cells[1], [subtitle], center=True, size_pt=16)
    _write_cell_lines(table.rows[2].cells[0], [org_line1, org_line2], size_pt=12)
    _write_cell_lines(
        table.rows[2].cells[1],
        [issue_label, date_range],
        center=True,
        size_pt=12,
    )


def _format_source_line(article: dict) -> str:
    parts: list[str] = []
    for entry in article.get("source_urls", []):
        name = entry.get("name", "").strip()
        url = entry.get("url", "").strip()
        if name and url:
            parts.append(f"{name} {url}")
        elif url:
            parts.append(url)
    if parts:
        return f"（信息来源：{'；'.join(parts)}）"
    legacy = article.get("source", "")
    legacy_url = article.get("url", "")
    if legacy and legacy_url:
        return f"（信息来源：{legacy} {legacy_url}）"
    if legacy:
        return f"（信息来源：{legacy}）"
    return ""


def _validate_data(data: dict) -> list[str]:
    warnings: list[str] = []
    summary = data.get("summary", [])
    articles = data.get("articles", [])
    if len(summary) != ISSUE_ARTICLE_COUNT:
        warnings.append(f"摘要应为 {ISSUE_ARTICLE_COUNT} 条，当前 {len(summary)} 条")
    if len(articles) != ISSUE_ARTICLE_COUNT:
        warnings.append(f"正文应为 {ISSUE_ARTICLE_COUNT} 条，当前 {len(articles)} 条")
    for i, art in enumerate(articles, 1):
        body = "".join(art.get("paragraphs", []))
        chars = count_chinese_chars(body)
        if chars > MAX_BODY_CHARS:
            warnings.append(f"第 {i} 条（{art.get('region')}）正文 {chars} 字，超过 {MAX_BODY_CHARS} 字")
        if not art.get("source_urls"):
            warnings.append(f"第 {i} 条缺少 source_urls")
    return warnings


def build_document(data: dict, template: Path | None = None) -> Document:
    meta = data["meta"]
    template_path = template or DEFAULT_TEMPLATE
    if not template_path.is_file():
        raise FileNotFoundError(f"Word 模板不存在: {template_path}")

    doc = Document(template_path)
    _clear_body_keep_header_table(doc)
    _fill_header_table(doc, meta)

    _add_run_paragraph(
        doc,
        "【本期摘要】",
        style="Normal Indent",
        east_asia=FONT_SECTION,
        bold=True,
        first_line_indent=0,
    )
    for item in data["summary"]:
        _add_run_paragraph(
            doc,
            f"{item['region']}：{item['title']}",
            style="List Paragraph",
            east_asia=FONT_BODY,
            line_spacing=LINE_SPACING_BODY,
        )

    _add_run_paragraph(
        doc,
        "【主要内容】",
        style="Normal Indent",
        east_asia=FONT_SECTION,
        bold=True,
        first_line_indent=0,
    )

    for article in data["articles"]:
        _add_run_paragraph(
            doc,
            f"{article['region']}：{article['title']}",
            style="Normal",
            east_asia=FONT_SECTION,
            bold=True,
            left_indent=INDENT_TITLE_LEFT,
            first_line_indent=0,
            line_spacing=LINE_SPACING_BODY,
        )

        for para_text in article["paragraphs"]:
            _add_run_paragraph(
                doc,
                para_text,
                style="Normal Indent",
                east_asia=FONT_BODY,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=INDENT_BODY_FIRST,
                line_spacing=LINE_SPACING_BODY,
            )

        source_line = _format_source_line(article)
        if source_line:
            _add_run_paragraph(
                doc,
                source_line,
                style="Normal Indent",
                east_asia=FONT_BODY,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
                line_spacing=LINE_SPACING_BODY,
            )

        _add_run_paragraph(doc, "", style="Normal Indent", first_line_indent=0)

    editors = meta.get("editors", [])
    if editors:
        _add_run_paragraph(
            doc,
            f"本期编辑：{'，'.join(editors)}",
            style="Normal",
            east_asia=FONT_BODY,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=INDENT_BODY_FIRST,
            line_spacing=LINE_SPACING_BODY,
        )

    return doc


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="生成海外油气投资环境双周报 Word")
    parser.add_argument("--input", "-i", required=True, help="结构化 JSON 路径")
    parser.add_argument("--output", "-o", required=True, help="输出 .docx 路径")
    parser.add_argument(
        "--template",
        "-t",
        default=str(DEFAULT_TEMPLATE),
        help="Word 模板（默认 references/sample_issue60.docx）",
    )
    args = parser.parse_args(argv)

    input_path = Path(args.input)
    output_path = Path(args.output)

    with input_path.open(encoding="utf-8") as f:
        data = json.load(f)

    for key in ("meta", "summary", "articles"):
        if key not in data:
            print(f"JSON 缺少必填字段: {key}", file=sys.stderr)
            return 1

    for warn in _validate_data(data):
        print(f"警告: {warn}", file=sys.stderr)

    doc = build_document(data, template=Path(args.template))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"已生成: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
